import io
import os
from typing import Any, Dict, List

import PyPDF2
from docx import Document
from fastapi import UploadFile
from langchain_text_splitters import CharacterTextSplitter

from src.config import settings
from src.http.exceptions.docs_excpetions import InvalidFormatException
from src.lib.langchain import LangChainClient
from src.schemas.docs_schema import DocsIndexingRequest, DocsType


class DocsService:
    def __init__(self, db_file: str = settings.path_db_file):
        self.llm_client = LangChainClient(db_file=db_file)
        self.valid_extensions = {
            ".pdf": DocsType.PDF,
            ".docx": DocsType.DOCX,
            ".txt": DocsType.TXT,
            ".html": DocsType.HTML,
            ".htm": DocsType.HTML,
            ".md": DocsType.MD,
            ".markdown": DocsType.MD,
        }

    def _get_file_format(self, file: UploadFile) -> DocsType:
        if not file.filename:
            raise InvalidFormatException("Nome do arquivo não encontrado")

        _, file_extension = os.path.splitext(file.filename)
        file_extension = file_extension.lower()

        if file_extension not in self.valid_extensions:
            raise InvalidFormatException(file_extension)

        return self.valid_extensions[file_extension]

    def _extract_text_from_pdf(self, content: bytes) -> str:
        try:
            pdf_file = io.BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            if len(pdf_reader.pages) == 0:
                raise Exception("PDF está vazio ou corrompido")

            text_parts = []
            for page_num, page in enumerate(pdf_reader.pages, 1):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_parts.append(f"--- Página {page_num} ---\n{page_text}")
                except Exception as e:
                    text_parts.append(
                        f"--- Página {page_num} ---\n[Erro ao extrair texto: {str(e)}]"
                    )

            if not text_parts:
                raise Exception(
                    "Não foi possível extrair texto de nenhuma página do PDF"
                )

            full_text = "\n\n".join(text_parts)
            return full_text.strip()

        except Exception as e:
            raise Exception(f"Erro ao processar PDF: {str(e)}")

    def _extract_text_from_docx(self, content: bytes) -> str:
        try:
            docx_file = io.BytesIO(content)
            doc = Document(docx_file)

            text_parts = []

            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))

            if not text_parts:
                raise Exception("Documento DOCX está vazio")

            full_text = "\n\n".join(text_parts)
            return full_text.strip()

        except Exception as e:
            raise Exception(f"Erro ao processar DOCX: {str(e)}")

    def _extract_text_from_file(
        self, content: bytes, file_format: DocsType, filename: str
    ) -> str:
        try:
            if file_format == DocsType.PDF:
                return self._extract_text_from_pdf(content)
            elif file_format == DocsType.DOCX:
                return self._extract_text_from_docx(content)
            elif file_format in [DocsType.TXT, DocsType.HTML, DocsType.MD]:
                encodings = ["utf-8", "latin-1", "cp1252", "iso-8859-1"]
                for encoding in encodings:
                    try:
                        text = content.decode(encoding)
                        if len(text.strip()) > 0 and not any(
                            char in text for char in ["\ufffd", "\x00"]
                        ):
                            return text
                    except (UnicodeDecodeError, UnicodeError):
                        continue

                return content.decode("utf-8", errors="ignore")
            else:
                raise Exception(f"Tipo de arquivo não suportado: {file_format}")

        except Exception as e:
            raise Exception(f"Erro ao processar arquivo {filename}: {str(e)}")

    async def indexa_documento(self, file: UploadFile, request: DocsIndexingRequest):
        file_format = self._get_file_format(file)

        content = await file.read()

        if len(content) == 0:
            raise Exception("Arquivo está vazio")

        try:
            text = self._extract_text_from_file(content, file_format, file.filename)

            if not text.strip():
                raise Exception("Não foi possível extrair conteúdo legível do arquivo")

            text_preview = text[:200].replace("\n", " ")
            print(f"Texto extraído de {file.filename}: {text_preview}...")

        except Exception as e:
            raise Exception(f"Erro ao processar {file.filename}: {str(e)}")

        chunk_size = 1500 if file_format == DocsType.PDF else 1000
        splitter = CharacterTextSplitter(
            chunk_size=chunk_size, chunk_overlap=150, separator="\n\n"
        )
        docs = splitter.split_text(text)

        if not docs:
            raise Exception("Não foi possível criar chunks do documento")

        enriched_texts = []
        for i, chunk in enumerate(docs):
            metadata_text = f"""Arquivo: {file.filename}
Nome: {request.name}
Descrição: {request.description}
Tipo: {file_format.value}
Chunk: {i+1}/{len(docs)}
Caracteres: {len(chunk)}

{chunk}"""
            enriched_texts.append(metadata_text)

        self.llm_client.add_texts(enriched_texts)

        return {
            "message": "Documento indexado com sucesso",
            "filename": file.filename,
            "file_type": file_format.value,
            "chunks_created": len(docs),
            "characters_processed": len(text),
            "chunk_size_used": chunk_size,
            "preview": text[:300] + "..." if len(text) > 300 else text,
        }

    async def search_docs(self, query: str, limit: int = 5) -> List[Dict[str, Any]]:
        if not query.strip():
            return []

        results = self.llm_client.similarity_search(query, k=limit)

        if not results:
            return []

        formatted_results = []
        for i, doc in enumerate(results):
            content = doc.page_content
            lines = content.split("\n")

            filename = ""
            doc_name = ""
            doc_type = ""
            chunk_info = ""

            for line in lines[:8]:
                if line.startswith("Arquivo:"):
                    filename = line.replace("Arquivo:", "").strip()
                elif line.startswith("Nome:"):
                    doc_name = line.replace("Nome:", "").strip()
                elif line.startswith("Tipo:"):
                    doc_type = line.replace("Tipo:", "").strip()
                elif line.startswith("Chunk:"):
                    chunk_info = line.replace("Chunk:", "").strip()

            content_start = 0
            for j, line in enumerate(lines):
                if line.strip() == "" and j > 5:
                    content_start = j + 1
                    break

            clean_content = (
                "\n".join(lines[content_start:]) if content_start > 0 else content
            )

            preview_length = 300 if doc_type == "PDF" else 200
            preview = clean_content[:preview_length]
            if len(clean_content) > preview_length:
                last_space = preview.rfind(" ")
                if last_space > preview_length * 0.8:
                    preview = preview[:last_space]
                preview += "..."

            formatted_results.append(
                {
                    "rank": i + 1,
                    "content": clean_content,
                    "filename": filename,
                    "document_name": doc_name,
                    "document_type": doc_type,
                    "chunk": chunk_info,
                    "preview": preview,
                    "content_length": len(clean_content),
                    "metadata": getattr(doc, "metadata", {}),
                }
            )

        return formatted_results

    async def search_docs_with_context(
        self, query: str, limit: int = 5
    ) -> Dict[str, Any]:
        results = await self.search_docs(query, limit)

        if not results:
            return {"query": query, "found_documents": 0, "results": [], "context": ""}

        context_parts = []
        for result in results:
            context_parts.append(
                f"[{result['document_name']} - Chunk {result['chunk']}]: {result['content']}"
            )

        context = "\n\n---\n\n".join(context_parts)

        total_chars = len(context)
        estimated_tokens = total_chars // 4

        return {
            "query": query,
            "found_documents": len(results),
            "results": results,
            "context": context,
            "context_stats": {
                "total_characters": total_chars,
                "estimated_tokens": estimated_tokens,
                "chunks_included": len(results),
            },
        }
