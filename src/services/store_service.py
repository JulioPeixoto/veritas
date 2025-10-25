import io
import os
from typing import List

import PyPDF2
from docx import Document
from fastapi import UploadFile
from langchain_text_splitters import CharacterTextSplitter

from src.api.exceptions.store_excpetions import InvalidFormatException
from src.config import settings
from src.lib.clients.langchain import LangChainClient
from src.schemas.store_schema import (
    ContextStats,
    DocsIndexingResponse,
    DocsType,
    DocumentIndexResult,
    SearchDocResult,
    SearchDocsResponse,
    SearchDocsWithContextResponse,
)


class StoreService:
    def __init__(self, db_file: str = settings.path_db_file):
        self.llm_client = LangChainClient(db_file=db_file)
        self.valid_extensions = {
            ".pdf": DocsType.PDF,
            ".docx": DocsType.DOCX,
            ".txt": DocsType.TXT,
            ".html": DocsType.HTML,
            ".htm": DocsType.HTML,
            ".md": DocsType.MD,
            ".markdown": DocsType.MD,
        }

    def _get_file_format(self, file: UploadFile) -> DocsType:
        if not file.filename:
            raise InvalidFormatException("Nome do arquivo não encontrado")

        _, file_extension = os.path.splitext(file.filename)
        file_extension = file_extension.lower()

        if file_extension not in self.valid_extensions:
            raise InvalidFormatException(file_extension)

        return self.valid_extensions[file_extension]

    def _get_document_name(self, filename: str) -> str:
        name_without_ext = os.path.splitext(filename)[0]
        clean_name = name_without_ext.replace("_", " ").replace("-", " ")
        return clean_name.title().strip()

    def _extract_text_from_pdf(self, content: bytes) -> str:
        try:
            pdf_file = io.BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            if len(pdf_reader.pages) == 0:
                raise Exception("PDF está vazio ou corrompido")

            text_parts = []
            for page_num, page in enumerate(pdf_reader.pages, 1):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_parts.append(f"--- Página {page_num} ---\n{page_text}")
                except Exception as e:
                    text_parts.append(
                        f"--- Página {page_num} ---\n[Erro ao extrair texto: {str(e)}]"
                    )

            if not text_parts:
                raise Exception(
                    "Não foi possível extrair texto de nenhuma página do PDF"
                )

            full_text = "\n\n".join(text_parts)
            return full_text.strip()

        except Exception as e:
            raise Exception(f"Erro ao processar PDF: {str(e)}")

    def _extract_text_from_docx(self, content: bytes) -> str:
        try:
            docx_file = io.BytesIO(content)
            doc = Document(docx_file)

            text_parts = []

            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))

            if not text_parts:
                raise Exception("Documento DOCX está vazio")

            full_text = "\n\n".join(text_parts)
            return full_text.strip()

        except Exception as e:
            raise Exception(f"Erro ao processar DOCX: {str(e)}")

    def _extract_text_from_file(
        self, content: bytes, file_format: DocsType, filename: str
    ) -> str:
        try:
            if file_format == DocsType.PDF:
                return self._extract_text_from_pdf(content)
            elif file_format == DocsType.DOCX:
                return self._extract_text_from_docx(content)
            elif file_format in [DocsType.TXT, DocsType.HTML, DocsType.MD]:
                encodings = ["utf-8", "latin-1", "cp1252", "iso-8859-1"]
                for encoding in encodings:
                    try:
                        text = content.decode(encoding)
                        if len(text.strip()) > 0 and not any(
                            char in text for char in ["\ufffd", "\x00"]
                        ):
                            return text
                    except (UnicodeDecodeError, UnicodeError):
                        continue

                return content.decode("utf-8", errors="ignore")
            else:
                raise Exception(f"Tipo de arquivo não suportado: {file_format}")

        except Exception as e:
            raise Exception(f"Erro ao processar arquivo {filename}: {str(e)}")

    async def indexa_documento(self, file: UploadFile) -> DocumentIndexResult:
        file_format = self._get_file_format(file)
        document_name = self._get_document_name(file.filename)

        content = await file.read()

        if len(content) == 0:
            raise Exception("Arquivo está vazio")

        try:
            text = self._extract_text_from_file(content, file_format, file.filename)

            if not text.strip():
                raise Exception("Não foi possível extrair conteúdo legível do arquivo")

        except Exception as e:
            raise Exception(f"Erro ao processar {file.filename}: {str(e)}")

        chunk_size = 1500 if file_format == DocsType.PDF else 1000
        splitter = CharacterTextSplitter(
            chunk_size=chunk_size, chunk_overlap=150, separator="\n\n"
        )
        docs = splitter.split_text(text)

        if not docs:
            raise Exception("Não foi possível criar chunks do documento")

        enriched_texts = []
        for i, chunk in enumerate(docs):
            metadata_text = f"""Arquivo: {file.filename}
Nome: {document_name}
Tipo: {file_format.value}
Chunk: {i+1}/{len(docs)}
Caracteres: {len(chunk)}

{chunk}"""
            enriched_texts.append(metadata_text)

        self.llm_client.add_texts(enriched_texts)

        return DocumentIndexResult(
            message="Documento indexado com sucesso",
            filename=file.filename,
            document_name=document_name,
            file_type=file_format.value,
            chunks_created=len(docs),
            characters_processed=len(text),
            chunk_size_used=chunk_size,
            preview=text[:300] + "..." if len(text) > 300 else text,
        )

    async def indexa_documentos(self, files: List[UploadFile]) -> DocsIndexingResponse:
        if not files:
            raise Exception("Nenhum arquivo fornecido")

        results = []
        total_chunks = 0
        total_characters = 0
        processed_files = 0
        errors = []

        for file in files:
            try:
                result = await self.indexa_documento(file)
                results.append(result)
                total_chunks += result.chunks_created
                total_characters += result.characters_processed
                processed_files += 1

            except Exception as e:
                error_msg = f"Erro ao processar {file.filename}: {str(e)}"
                errors.append(error_msg)
                print(error_msg)
                continue

        if processed_files == 0:
            raise Exception(
                f"Nenhum arquivo foi processado com sucesso. Erros: {'; '.join(errors)}"
            )

        return DocsIndexingResponse(
            processed_files=processed_files,
            results=results,
            total_chunks=total_chunks,
            total_characters=total_characters,
            errors=errors if errors else None,
        )

    async def search_docs(self, query: str, limit: int = 5) -> SearchDocsResponse:
        if not query.strip():
            return SearchDocsResponse(results=[])

        results = self.llm_client.similarity_search(query, k=limit)

        if not results:
            return SearchDocsResponse(results=[])

        formatted_results = []
        for i, doc in enumerate(results):
            content = doc.page_content
            lines = content.split("\n")

            filename = ""
            doc_name = ""
            doc_type = ""
            chunk_info = ""

            for line in lines[:7]:
                if line.startswith("Arquivo:"):
                    filename = line.replace("Arquivo:", "").strip()
                elif line.startswith("Nome:"):
                    doc_name = line.replace("Nome:", "").strip()
                elif line.startswith("Tipo:"):
                    doc_type = line.replace("Tipo:", "").strip()
                elif line.startswith("Chunk:"):
                    chunk_info = line.replace("Chunk:", "").strip()

            content_start = 0
            for j, line in enumerate(lines):
                if line.strip() == "" and j > 4:
                    content_start = j + 1
                    break

            clean_content = (
                "\n".join(lines[content_start:]) if content_start > 0 else content
            )

            preview_length = 300 if doc_type == "PDF" else 200
            preview = clean_content[:preview_length]
            if len(clean_content) > preview_length:
                last_space = preview.rfind(" ")
                if last_space > preview_length * 0.8:
                    preview = preview[:last_space]
                preview += "..."

            formatted_results.append(
                SearchDocResult(
                    rank=i + 1,
                    content=clean_content,
                    filename=filename,
                    document_name=doc_name,
                    document_type=doc_type,
                    chunk=chunk_info,
                    preview=preview,
                    content_length=len(clean_content),
                    metadata=getattr(doc, "metadata", {}),
                )
            )

        return SearchDocsResponse(results=formatted_results)

    async def search_docs_with_context(
        self, query: str, limit: int = 5
    ) -> SearchDocsWithContextResponse:
        results = await self.search_docs(query, limit)

        if not results.results:
            return SearchDocsWithContextResponse(
                query=query,
                found_documents=0,
                results=[],
                context="",
                context_stats=ContextStats(
                    total_characters=0,
                    estimated_tokens=0,
                    chunks_included=0,
                ),
            )

        context_parts = []
        for result in results.results:
            context_parts.append(
                f"[{result.document_name} - Chunk {result.chunk}]: {result.content}"
            )

        context = "\n\n---\n\n".join(context_parts)

        total_chars = len(context)
        estimated_tokens = total_chars // 4

        return SearchDocsWithContextResponse(
            query=query,
            found_documents=len(results.results),
            results=results.results,
            context=context,
            context_stats=ContextStats(
                total_characters=total_chars,
                estimated_tokens=estimated_tokens,
                chunks_included=len(results.results),
            ),
        )
